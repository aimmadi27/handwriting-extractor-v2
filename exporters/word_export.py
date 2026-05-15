import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_word(page_results: dict) -> bytes:
    doc = Document()

    # Compact margins
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    first_page = True
    for page_num in sorted(page_results.keys()):
        page = page_results[page_num]

        if not first_page:
            doc.add_page_break()
        first_page = False

        # Page title
        h = doc.add_heading(page.get("title", f"Page {page_num}"), level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        doc_type = page.get("doc_type", "")
        if doc_type:
            p = doc.add_paragraph(doc_type.upper())
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        for section in page.get("sections", []):
            stype = section.get("type")
            title = section.get("title")

            if title:
                doc.add_heading(title, level=2)

            if stype == "key_value":
                pairs = section.get("pairs") or []
                if pairs:
                    table = doc.add_table(rows=0, cols=2)
                    table.style = "Table Grid"
                    for pair in pairs:
                        row = table.add_row()
                        row.cells[0].text = str(pair.get("key", ""))
                        row.cells[0].paragraphs[0].runs[0].bold = True
                        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                        row.cells[1].text = str(pair.get("value", "") or "")
                    doc.add_paragraph()

            elif stype == "table":
                columns = section.get("columns") or []
                rows = section.get("rows") or []
                if columns:
                    table = doc.add_table(rows=1, cols=len(columns))
                    table.style = "Table Grid"
                    for i, col in enumerate(columns):
                        cell = table.rows[0].cells[i]
                        cell.text = str(col)
                        cell.paragraphs[0].runs[0].bold = True
                    for row in rows:
                        r = table.add_row()
                        for i, val in enumerate(row[:len(columns)]):
                            r.cells[i].text = str(val or "")
                    doc.add_paragraph()

            elif stype == "qa_pair":
                for item in (section.get("items") or []):
                    q_para = doc.add_paragraph()
                    q_run = q_para.add_run(str(item.get("question", "")))
                    q_run.bold = True
                    q_run.font.size = Pt(11)

                    answer = str(item.get("answer", "") or "")
                    a_para = doc.add_paragraph(answer)
                    a_para.paragraph_format.left_indent = Pt(16)
                    doc.add_paragraph()

            elif stype == "paragraph":
                text = str(section.get("text", "") or "")
                for block in text.split("\n\n"):
                    if block.strip():
                        doc.add_paragraph(block)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
